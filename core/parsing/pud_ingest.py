"""Извлечение структуры оценивания из ФАЙЛА ПУД (контур 1, первый сценарий).

Преподаватель присылает свой ПУД файлом (PDF/DOCX/HTML из конструктора dp.hse.ru).
Формат «Системы оценивания» в конструкторе стандартизирован: есть строка формулы
вида «Активность * 0.1 + Блиц * 0.2 + Контрольная работа * 0.5». Её и вытаскиваем
детерминированно (без LLM). Qwen — запасной путь для нестандартных ПУД (позже).
"""
from __future__ import annotations

import re
from pathlib import Path

# Формула: последовательность «<название> * <вес>», соединённая знаком +.
_FORMULA_RE = re.compile(
    r"([^\s*+][^*+]{1,80}?\s*\*\s*\d+(?:[.,]\d+)?"
    r"(?:\s*\+\s*[^*+]{1,80}?\s*\*\s*\d+(?:[.,]\d+)?)+)"
)


def extract_text(path: str) -> str:
    """Достаёт текст из файла ПУД по расширению. Поддержка: html, docx, pdf, txt."""
    p = Path(path)
    ext = p.suffix.lower()
    if ext in (".html", ".htm"):
        raw = p.read_text(encoding="utf-8", errors="ignore")
        raw = re.sub(r"(?s)<(script|style).*?</\1>", " ", raw)
        return re.sub(r"<[^>]+>", "\n", raw)
    if ext == ".docx":
        import docx  # python-docx
        d = docx.Document(str(p))
        parts = [par.text for par in d.paragraphs]
        for tbl in d.tables:
            for row in tbl.rows:
                parts.append(" ".join(c.text for c in row.cells))
        return "\n".join(parts)
    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(p))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if ext in (".txt", ".md"):
        return p.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Неподдерживаемый формат ПУД: {ext} (нужен PDF, DOCX, HTML или TXT)")


def find_formula(text: str) -> str | None:
    """Находит строку формулы оценивания в тексте ПУД.
    Приоритет — секция после метки «Формула оценивания:» (так первый элемент не
    склеивается с преамбулой); иначе — поиск паттерна по всему тексту."""
    flat = re.sub(r"\s+", " ", text)
    label = re.search(r"Формула\s+оценивани[ея]\s*:?", flat)
    search_space = flat[label.end():] if label else flat
    m = _FORMULA_RE.search(search_space)
    if not m:
        return None
    return m.group(1).strip()


def find_title(text: str) -> str:
    """Best-effort: название дисциплины (строка после метки «Название»)."""
    lines = [l.strip() for l in text.splitlines()]
    for i, l in enumerate(lines):
        if l == "Название" and i + 1 < len(lines) and lines[i + 1]:
            return lines[i + 1][:150]
    return "Ведомость по ПУД"
